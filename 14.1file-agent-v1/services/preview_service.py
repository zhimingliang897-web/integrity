import os
import io
import base64
from pathlib import Path
from typing import Optional, Tuple
from fastapi.responses import StreamingResponse, FileResponse

from app.config import settings


class PreviewService:
    def __init__(self):
        self.root_path = Path(settings.root_path)
        # 允许预览的根目录：主目录 + 所有挂载点
        self.allowed_roots = [self.root_path.resolve()]
        for mount in settings.mounts:
            mount_path = mount.get("path")
            if mount_path:
                try:
                    p = Path(mount_path)
                    if p.exists():
                        self.allowed_roots.append(p.resolve())
                except Exception:
                    continue
        
        self.image_exts = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.svg', '.ico'}
        self.video_exts = {'.mp4', '.webm', '.ogg', '.mov'}
        self.audio_exts = {'.mp3', '.wav', '.flac', '.aac', '.ogg', '.m4a'}
        self.pdf_exts = {'.pdf'}
        self.text_exts = {'.txt', '.md', '.py', '.js', '.ts', '.html', '.css', '.json', '.xml', '.log', '.csv'}
        self.doc_exts = {'.doc', '.docx'}
    
    def _is_path_allowed(self, path: str) -> bool:
        """检查路径是否在允许预览的根目录（主目录或挂载目录）之下。"""
        try:
            abs_path = Path(path).resolve()
            for root in self.allowed_roots:
                if str(abs_path).startswith(str(root)):
                    return True
            return False
        except Exception:
            return False
    
    def get_preview_type(self, ext: str) -> str:
        ext = ext.lower()
        if ext in self.image_exts:
            return "image"
        elif ext in self.video_exts:
            return "video"
        elif ext in self.audio_exts:
            return "audio"
        elif ext in self.pdf_exts:
            return "pdf"
        elif ext in self.text_exts or ext in self.doc_exts:
            return "text"
        else:
            return "unknown"
    
    def can_preview(self, path: str) -> bool:
        ext = Path(path).suffix.lower()
        return self.get_preview_type(ext) != "unknown"
    
    def get_file_response(self, path: str, preview: bool = False) -> FileResponse:
        file_path = Path(path)
        
        if not self._is_path_allowed(path):
            raise ValueError("路径不允许访问")
        
        if not file_path.exists():
            raise FileNotFoundError("文件不存在")
        
        return FileResponse(
            path=str(file_path),
            filename=file_path.name,
            media_type=self._get_mime_type(path)
        )
    
    def get_text_preview(self, path: str, max_size: int = 100000) -> str:
        file_path = Path(path)
        
        if not self._is_path_allowed(path):
            raise ValueError("路径不允许访问")
        
        if not file_path.exists():
            raise FileNotFoundError("文件不存在")
        
        ext = file_path.suffix.lower()
        
        if ext in {'.doc', '.docx'}:
            return self.get_docx_text(path, max_size)
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read(max_size)
        except Exception as e:
            return f"无法读取文件: {str(e)}"
    
    def get_docx_text(self, path: str, max_size: int = 100000) -> str:
        file_path = Path(path)
        
        if not self._is_path_allowed(path):
            raise ValueError("路径不允许访问")
        
        if not file_path.exists():
            raise FileNotFoundError("文件不存在")
        
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            text_parts = []
            total_len = 0
            
            for para in doc.paragraphs:
                if total_len >= max_size:
                    break
                text = para.text
                text_parts.append(text)
                total_len += len(text)
            
            for table in doc.tables:
                if total_len >= max_size:
                    break
                for row in table.rows:
                    for cell in row.cells:
                        if total_len >= max_size:
                            break
                        text_parts.append(cell.text)
                        total_len += len(cell.text)
            
            result = '\n'.join(text_parts)
            if len(result) > max_size:
                result = result[:max_size] + '\n... (内容过长，已截断)'
            
            return result
        except ImportError:
            return "无法预览：python-docx 库未安装"
        except Exception as e:
            return f"无法读取Word文档: {str(e)}"
    
    def get_image_base64(self, path: str, max_size: int = 10 * 1024 * 1024) -> str:
        file_path = Path(path)
        
        if not self._is_path_allowed(path):
            raise ValueError("路径不允许访问")
        
        if not file_path.exists():
            raise FileNotFoundError("文件不存在")
        
        file_size = file_path.stat().st_size
        if file_size > max_size:
            raise ValueError(f"文件过大 ({file_size // (1024*1024)}MB)，请下载查看")
        
        with open(file_path, 'rb') as f:
            img_data = f.read()
        
        ext = file_path.suffix.lower()
        mime_map = {
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif',
            '.webp': 'image/webp',
            '.bmp': 'image/bmp',
            '.svg': 'image/svg+xml',
            '.ico': 'image/x-icon'
        }
        mime_type = mime_map.get(ext, 'image/jpeg')
        
        return f"data:{mime_type};base64,{base64.b64encode(img_data).decode('utf-8')}"
    
    def get_image_thumbnail(self, path: str, size: Tuple[int, int] = (200, 200)):
        try:
            from PIL import Image
            
            file_path = Path(path)
            
            if not self._is_path_allowed(path):
                raise ValueError("路径不允许访问")
            
            if not file_path.exists():
                raise FileNotFoundError("文件不存在")
            
            img = Image.open(file_path)
            img.thumbnail(size)
            
            img_byte_arr = io.BytesIO()
            img_format = img.format or 'JPEG'
            if img_format == 'PNG':
                img.save(img_byte_arr, format='PNG')
            else:
                if img.mode in ('RGBA', 'P'):
                    img = img.convert('RGB')
                img.save(img_byte_arr, format='JPEG')
            
            img_byte_arr.seek(0)
            return img_byte_arr
        except ImportError:
            raise RuntimeError("Pillow 库未安装，无法生成缩略图")
        except Exception as e:
            raise RuntimeError(f"生成缩略图失败: {str(e)}")
    
    def _get_mime_type(self, path: str) -> str:
        import mimetypes
        mime_type, _ = mimetypes.guess_type(path)
        return mime_type or 'application/octet-stream'